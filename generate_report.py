from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = "/Users/atharvakulkarni/Desktop/MS SE SJSU/CMPE-282/HW5 - Security/Screenshots"
GITHUB_URL = "https://github.com/AK9175/cmpe282-hw5-security"

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_screenshot(filename, caption):
    path = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.8))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

# ─────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_heading('CMPE-282: Cloud Services\nHomework #5 – Security', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\nStudent: Atharva Kulkarni\nEmail: atharva.kulkarni@sjsu.edu\nDate: April 25, 2026\n')
info.add_run(f'\nGitHub Repository: {GITHUB_URL}').bold = True

doc.add_page_break()

# ─────────────────────────────────────────────
# PART A
# ─────────────────────────────────────────────
heading1('Part A: Configure Apache Web Server with SSL/TLS')
doc.add_paragraph(
    'This section documents the configuration of an Apache Web Server on Amazon Linux 2 '
    'to use SSL/TLS, verified via deep analysis on SSL Labs (ssllabs.com).'
)

# Step A1
heading2('Step A.1 – Launch EC2 Instance')
doc.add_paragraph(
    'An Amazon Linux 2 t2.micro EC2 instance named hw5-ssl-server was launched using the AWS CLI. '
    'A new security group was created allowing inbound traffic on ports 22 (SSH), 80 (HTTP), and 443 (HTTPS).'
)
add_code('aws ec2 run-instances --image-id ami-0058102e678cd65ab --instance-type t2.micro \\\n'
         '  --key-name hw5-key --security-group-ids sg-03b91fea232ceb60a \\\n'
         '  --tag-specifications \'ResourceType=instance,Tags=[{Key=Name,Value=hw5-ssl-server}]\'')
doc.add_paragraph()
add_screenshot('SS1.png', 'Figure 1: EC2 instance hw5-ssl-server in Running state (Public IP: 54.219.251.209)')

# Step A2
heading2('Step A.2 – Install Apache and mod_ssl')
doc.add_paragraph(
    'Apache HTTP Server and mod_ssl were installed on the EC2 instance via SSH. '
    'Apache was started and enabled to launch on system boot.'
)
add_code('sudo yum install -y httpd mod_ssl\nsudo systemctl start httpd\nsudo systemctl enable httpd')
doc.add_paragraph()
add_screenshot('SS2.png', 'Figure 2: Apache HTTP Server active and running on EC2')

# Step A3
heading2('Step A.3 – Obtain SSL Certificate via Let\'s Encrypt (Certbot)')
doc.add_paragraph(
    'Certbot was installed via EPEL and used to obtain a trusted SSL/TLS certificate from Let\'s Encrypt '
    'for the domain 54.219.251.209.nip.io (nip.io provides free wildcard DNS for any IP, eliminating '
    'the need to purchase a domain). The certificate was automatically configured in Apache.'
)
add_code('sudo amazon-linux-extras install epel -y\n'
         'sudo yum install -y certbot python2-certbot-apache\n'
         'sudo certbot --apache -d 54.219.251.209.nip.io --non-interactive --agree-tos -m atharva.kulkarni@sjsu.edu')
doc.add_paragraph(
    'HSTS (HTTP Strict Transport Security) was also enabled to enforce HTTPS connections, '
    'improving the SSL Labs score.'
)
add_code('Header always set Strict-Transport-Security "max-age=63072000; includeSubDomains"')
doc.add_paragraph()
add_screenshot('SS3.png', 'Figure 3: HTTPS site loading with valid SSL certificate in browser')

# Step A4
heading2('Step A.4 – SSL Labs Deep Analysis')
doc.add_paragraph(
    'The SSL configuration was analyzed using the Qualys SSL Labs Server Test at '
    'https://www.ssllabs.com/ssltest/analyze.html. The server achieved an A- overall rating, '
    'with the certificate trusted by Mozilla, Apple, Android, Java, and Windows root stores. '
    'The certificate was issued by Let\'s Encrypt (R13 intermediate, ISRG Root X1).'
)
doc.add_paragraph('Key findings from SSL Labs analysis:')
items = [
    'Overall Grade: A-',
    'Certificate: Valid, trusted by all major root stores',
    'Protocol Support: TLS 1.2',
    'Key Exchange: Strong (ECDHE)',
    'Cipher Strength: Strong (AES-256)',
    'HSTS: Configured with long duration',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
add_screenshot('SS4.png', 'Figure 4: SSL Labs analysis showing A- grade for 54.219.251.209.nip.io')

doc.add_paragraph(
    f'Full setup script available on GitHub: {GITHUB_URL}/blob/main/apache-ssl-setup.sh'
)

doc.add_page_break()

# ─────────────────────────────────────────────
# PART B
# ─────────────────────────────────────────────
heading1('Part B: Use IAM Roles to Grant Access to an AWS Application on EC2')
doc.add_paragraph(
    'This section demonstrates using IAM Roles to securely grant an EC2 instance read-only '
    'access to Amazon S3 without embedding any credentials in the application code.'
)

# Step B1
heading2('Step B.1 – Create IAM Role with S3 Read-Only Access')
doc.add_paragraph(
    'An IAM role named hw5-s3-readonly-role was created with a trust policy allowing EC2 to assume it. '
    'The AWS managed policy AmazonS3ReadOnlyAccess was attached, and an instance profile was created '
    'to associate the role with an EC2 instance.'
)
add_code('aws iam create-role --role-name hw5-s3-readonly-role \\\n'
         '  --assume-role-policy-document \'{"Version":"2012-10-17","Statement":[{"Effect":"Allow",\n'
         '  "Principal":{"Service":"ec2.amazonaws.com"},"Action":"sts:AssumeRole"}]}\'\n\n'
         'aws iam attach-role-policy --role-name hw5-s3-readonly-role \\\n'
         '  --policy-arn arn:aws:iam::aws:policy/AmazonS3ReadOnlyAccess\n\n'
         'aws iam create-instance-profile --instance-profile-name hw5-s3-readonly-profile\n'
         'aws iam add-role-to-instance-profile --instance-profile-name hw5-s3-readonly-profile \\\n'
         '  --role-name hw5-s3-readonly-role')
doc.add_paragraph()
add_screenshot('BSS1.png', 'Figure 5: IAM role hw5-s3-readonly-role with AmazonS3ReadOnlyAccess policy attached')

# Step B2
heading2('Step B.2 – Launch EC2 Instance with IAM Role')
doc.add_paragraph(
    'A second EC2 instance named hw5-s3-app-server was launched with the IAM instance profile attached. '
    'This grants the instance automatic, temporary, rotating credentials to access AWS services — '
    'no hardcoded keys required.'
)
add_code('aws ec2 run-instances --image-id ami-0058102e678cd65ab --instance-type t2.micro \\\n'
         '  --key-name hw5-key --iam-instance-profile Name=hw5-s3-readonly-profile \\\n'
         '  --tag-specifications \'ResourceType=instance,Tags=[{Key=Name,Value=hw5-s3-app-server}]\'')
doc.add_paragraph()
add_screenshot('BSS2.png', 'Figure 6: EC2 instance hw5-s3-app-server with hw5-s3-readonly-role attached')

# Step B3
heading2('Step B.3 – Transfer Application to EC2')
doc.add_paragraph(
    'A Python application (s3_app.py) was written using the AWS SDK for Python (boto3). '
    'The application lists all S3 buckets and their objects using credentials automatically '
    'provided by the EC2 instance metadata service via the attached IAM role. '
    'The file was transferred to the EC2 instance using SCP.'
)
add_code(f'# Source code: {GITHUB_URL}/blob/main/s3_app.py\n\n'
         'scp -i hw5-key.pem s3_app.py ec2-user@3.101.191.21:~/s3_app.py')
doc.add_paragraph()
add_screenshot('BSS3.png', 'Figure 7: s3_app.py transferred to EC2 instance via SCP')

# Step B4
heading2('Step B.4 – Run Application on EC2')
doc.add_paragraph(
    'The application was executed on the EC2 instance. boto3 automatically retrieved temporary '
    'credentials from the EC2 Instance Metadata Service (IMDS) using the attached IAM role — '
    'no AWS keys were configured on the instance. The app successfully listed all S3 buckets '
    'and their contents, confirming read-only access.'
)
add_code('pip3 install boto3\npython3 ~/s3_app.py')
doc.add_paragraph()
add_screenshot('BSS4.png', 'Figure 8: s3_app.py running on EC2 — S3 buckets listed via IAM role')

# ─────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────
doc.add_page_break()
heading1('References')
refs = [
    ('AWS Documentation – SSL on Amazon Linux 2',
     'https://docs.aws.amazon.com/linux/al2/ug/SSL-on-amazon-linux-2.html'),
    ('Qualys SSL Labs Server Test',
     'https://www.ssllabs.com/ssltest/analyze.html'),
    ('AWS Documentation – IAM Roles for EC2',
     'https://docs.aws.amazon.com/AWSEC2/latest/UserGuide/iam-roles-for-amazon-ec2.html'),
    ('GitHub Repository – cmpe282-hw5-security',
     GITHUB_URL),
    ('s3_app.py – S3 SDK Application',
     f'{GITHUB_URL}/blob/main/s3_app.py'),
    ('apache-ssl-setup.sh – Apache SSL Setup Script',
     f'{GITHUB_URL}/blob/main/apache-ssl-setup.sh'),
    ('iam-setup.sh – IAM Role Setup Script',
     f'{GITHUB_URL}/blob/main/iam-setup.sh'),
]
for title, url in refs:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(url)

# Save
out = "/Users/atharvakulkarni/Desktop/MS SE SJSU/CMPE-282/HW5 - Security/HW5_Security_Atharva_Kulkarni.docx"
doc.save(out)
print(f"Word document saved: {out}")
